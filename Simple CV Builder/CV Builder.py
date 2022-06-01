from docx import Document
import os


os.chdir(os.path.dirname(__file__))

document = Document()

# name phone number and email details
name = input("What is your name?\n")
phone_number = input("What is your phone number?\n")
email = input("What is your email?\n")

document.add_paragraph(f"{name} | {phone_number} | {email}")


# about me
document.add_heading("About me")
about_me = input("Tell us about yourself?\n")
document.add_paragraph(about_me)

# work experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter company\n")
from_date = input("From Date\n")
to_date = input("To Date\n")

p.add_run(company + "\n").bold = True
p.add_run(f"{from_date} - {to_date} \n").italic = True


experience_details = input(f"Describe your experience at {company}\n")
p.add_run(experience_details)

while True:
    has_more_experiences = input("Do you have more experiences? Yes or No\n")
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()

        company = input("Enter company\n")
        from_date = input("From Date\n")
        to_date = input("To Date\n")

        p.add_run(company + " ").bold = True
        p.add_run(f"{from_date} - {to_date} \n").italic = True

        experience_details = input(f"Describe your experience at {company}\n")
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading("Skills")
skill = input("Enter skill\n")
p = document.add_paragraph(skill)
p.style = "List Bullet"

while True:
    has_more_skills = input("Do you have any more skills? Yes or No\n")
    if has_more_skills.lower == "yes":
        document.add_heading("Skills")
        skill = input("Enter skill\n")
        p = document.add_paragraph(skill)
        p.style = "List Bullet"
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV Generate by CV Builder"

document.save(f"{name} CV.docx")
